import os
import sqlite3
import subprocess
import tempfile
import shutil
from datetime import datetime
import pytz
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import (
    Updater, CommandHandler, CallbackQueryHandler, MessageHandler,
    Filters, ConversationHandler, CallbackContext
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request

# Flask-приложение для вебхуков
app = Flask(__name__)

# Токен бота и вебхук из переменных окружения
TOKEN = os.environ.get('BOT_TOKEN', '7511704960:AAFKDWgg2-cAzRxywX1gXK47OQRWJi72qGw')
WEBHOOK_URL = os.environ.get('WEBHOOK_URL', 'https://your-app.onrender.com/webhook')

# Состояния диалога
SELECT_TEMPLATE, INPUT_NAME, CHOOSE_DATE, INPUT_CUSTOM_DATE, ASK_SAVE = range(5)

# Часовой пояс Киева
kiev_tz = pytz.timezone('Europe/Kiev')

# Соответствие шаблонов
TEMPLATE_FILES = {
    'template_imperative': 'templates/template_imperative.docx',
    'template_ur': 'templates/template_ur.docx',
    'template_small_world': 'templates/template_small_world.docx',
}

def create_docx_template(filename, content):
    """Создаёт .docx файл с заданным содержимым."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    for item in content:
        if item['type'] == 'heading':
            p = doc.add_heading(item['text'], level=item.get('level', 1))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if item.get('center', False) else WD_ALIGN_PARAGRAPH.LEFT
        elif item['type'] == 'paragraph':
            p = doc.add_paragraph(item['text'])
            if item.get('bold', False):
                for run in p.runs:
                    run.bold = True
            if item.get('center', False):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif item['type'] == 'list':
            for text in item['items']:
                doc.add_paragraph(text, style='ListNumber' if item.get('numbered', True) else 'ListBullet')

    doc.save(filename)

def initialize_templates():
    """Создаёт .docx шаблоны, если они ещё не существуют."""
    os.makedirs('templates', exist_ok=True)

    # Шаблон Imperative
    imperative_content = [
        {'type': 'heading', 'text': 'IMPERATIVE PEOPLE LIMITED', 'center': True, 'level': 1},
        {'type': 'paragraph', 'text': 'Quinn Luke\nCompany number: 12463051\nUnit 15 Marston Business Park, Lower Hazeldines,\nMarston Moretaine, Bedfordshire, England, MK43 0XT\nhttps://imperative-people-limited.com/', 'center': True},
        {'type': 'heading', 'text': 'CONTRACT', 'center': True, 'level': 2},
        {'type': 'heading', 'text': '1. SUBJECT OF THE AGREEMENT', 'level': 3},
        {'type': 'paragraph', 'text': '1.1. Pursuant to this Agreement: Contractor - IMPERATIVE PEOPLE LIMITED\nCompany type - Private Limited Company number - 12463051\nRegistered office address: Unit 15 Marston Business Park, Lower Hazeldines, Marston Moretaine, Bedfordshire, England, MK43 0XT'},
        {'type': 'paragraph', 'text': 'Client:', 'bold': True},
        {'type': 'paragraph', 'text': 'The Contractor personally, at his own risk, provides the Client with services listed in paragraph 1.2 of this Agreement (hereinafter - "Services") within the period agreed by the Parties, and the Client accepts the Services provided by the Contractor and pays for the Services provided by him within the time, in the manner and in the amount established by this Agreement.'},
        {'type': 'list', 'items': ['1.2. Services provided by the Contractor to the Client in accordance with paragraph 1.1 of this Agreement:', '1.2.1. assistance and assistance in employment abroad.']},
        {'type': 'heading', 'text': '2. PROCEDURE FOR PERFORMANCE OF THE AGREEMENT', 'level': 3},
        {'type': 'paragraph', 'text': '2.1. The information required for the provision of Services under this Agreement is collected by the Contractor through its independent search, selection, systematization and analysis.'},
        {'type': 'paragraph', 'text': '2.2. Services are provided solely on the basis of information and documentation to be provided by the Client at the request and discretion of the Contractor, in accordance with the internal procedures of the latter.'},
        {'type': 'paragraph', 'text': '2.3. Term of providing Services:'},
        {'type': 'list', 'items': ['2.3.1. 30 (thirty) working days (excluding holidays or weekends) from the moment of payment in the amount and in the manner prescribed by paragraphs 4.1, 4.2 of this Agreement.']},
        # Добавьте остальные пункты по необходимости
        {'type': 'heading', 'text': 'SIGNATURES', 'level': 3},
        {'type': 'paragraph', 'text': 'THE CONTRACTOR:'},
        {'type': 'paragraph', 'text': 'Date:', 'bold': True},
        {'type': 'paragraph', 'text': 'THE CLIENT:'},
        {'type': 'paragraph', 'text': 'DATE:', 'bold': True},
        {'type': 'paragraph', 'text': 'SIGNATURE'},
    ]

    # Шаблон UR
    ur_content = [
        {'type': 'heading', 'text': 'UR RECRUITMENT LTD', 'center': True, 'level': 1},
        {'type': 'paragraph', 'text': 'RAFIQ Uziyan\nCompany number: 14593456\n38 Brockhurst Road, Birmingham, England, B36 8JB\nhttps://ur-recruitment.com/', 'center': True},
        {'type': 'heading', 'text': 'CONTRACT', 'center': True, 'level': 2},
        {'type': 'heading', 'text': 'SUBJECT OF THE AGREEMENT', 'level': 3},
        {'type': 'paragraph', 'text': '1.1. Pursuant to this Agreement:\nContractor - UR RECRUITMENT LTD\nCompany number 14593456, 38 Brockhurst Road, Birmingham, England, B36 8JB'},
        {'type': 'paragraph', 'text': 'Client:', 'bold': True},
        {'type': 'paragraph', 'text': 'The Contractor personally, at its own risk, provides the Client with services listed in paragraph 1.2 of this Agreement (hereinafter referred to as "Services") within the period agreed by the Parties. The Client accepts the Services provided by the Contractor and pays for the Services within the time, manner, and amount established by this Agreement.'},
        {'type': 'list', 'items': ['1.2. Services provided by the Contractor to the Client in accordance with paragraph 1.1 of this Agreement:', '1.2.1. Assistance in employment abroad.']},
        {'type': 'heading', 'text': 'Signatures:', 'level': 3},
        {'type': 'paragraph', 'text': 'Client'},
        {'type': 'paragraph', 'text': 'Date:', 'bold': True},
        {'type': 'paragraph', 'text': 'Contractor'},
        {'type': 'paragraph', 'text': 'Date:', 'bold': True},
    ]

    # Шаблон Small World
    small_world_content = [
        {'type': 'heading', 'text': 'SMALL WORLD RECRUITMENT', 'center': True, 'level': 1},
        {'type': 'paragraph', 'text': 'Company number: 05539195\nRedford & Co, 64 Baker Street, London, W1U 7GB\nTURNER Imogen Sarah', 'center': True},
        {'type': 'heading', 'text': 'CONTRACT', 'center': True, 'level': 2},
        {'type': 'heading', 'text': 'SUBJECT OF THE AGREEMENT', 'level': 3},
        {'type': 'paragraph', 'text': '1.1. Pursuant to this Agreement:\nContractor - SMALL WORLD RECRUITMENT LIMITED\nCompany number 05539195 Redford & Co, 64 Baker Street, London, W1U 7GB.'},
        {'type': 'paragraph', 'text': 'Client:', 'bold': True},
        {'type': 'paragraph', 'text': 'The Contractor personally, at its own risk, provides the Client with services listed in paragraph 1.2 of this Agreement (hereinafter referred to as "Services") within the period agreed by the Parties. The Client accepts the Services provided by the Contractor and pays for the Services within the time, manner, and amount established by this Agreement.'},
        {'type': 'list', 'items': ['1.2. Services provided by the Contractor to the Client in accordance with paragraph 1.1 of this Agreement:', '1.2.1. Assistance in employment abroad.']},
        {'type': 'heading', 'text': 'Signatures:', 'level': 3},
        {'type': 'paragraph', 'text': 'Client'},
        {'type': 'paragraph', 'text': 'Date:', 'bold': True},
        {'type': 'paragraph', 'text': 'Contractor'},
        {'type': 'paragraph', 'text': 'Date:', 'bold': True},
    ]

    for template, content in [
        ('templates/template_imperative.docx', imperative_content),
        ('templates/template_ur.docx', ur_content),
        ('templates/template_small_world.docx', small_world_content)
    ]:
        if not os.path.exists(template):
            create_docx_template(template, content)

def replace_text_in_paragraph(paragraph, key, value):
    """Замена текста в параграфе с сохранением форматирования."""
    if key in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if key in inline[i].text:
                inline[i].text = inline[i].text.replace(key, value)

def replace_text(doc, key, value):
    """Замена текста во всем документе Word."""
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, key, value)

def generate_and_send_document(update: Update, context: CallbackContext):
    """Генерация и отправка отредактированного документа в PDF."""
    template = context.user_data['template']
    client_name = context.user_data['client_name']
    date_time = context.user_data['date_time']
    chat_id = update.effective_chat.id

    context.bot.send_message(chat_id=chat_id, text="📄 Генерирую документ, подождите... ⏳")
    template_path = TEMPLATE_FILES[template]

    try:
        with tempfile.TemporaryDirectory() as tmpdirname:
            docx_path = os.path.join(tmpdirname, 'document.docx')
            pdf_path = os.path.join(tmpdirname, 'document.pdf')
            shutil.copy(template_path, docx_path)

            # Редактирование документа
            doc = Document(docx_path)
            replace_text(doc, "Client:", f"Client: {client_name}")
            replace_text(doc, "Date:", f"Date: {date_time}")
            replace_text(doc, "DATE:", f"DATE: {date_time}")
            doc.save(docx_path)

            # Конвертация в PDF
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', tmpdirname], check=True)

            # Отправка PDF
            with open(pdf_path, 'rb') as f:
                context.bot.send_document(chat_id=chat_id, document=f, caption="✅ Документ готов!")
    except Exception as e:
        context.bot.send_message(chat_id=chat_id, text=f"❌ Ошибка при создании документа: {str(e)}")

def start(update: Update, context: CallbackContext) -> None:
    """Приветственное сообщение."""
    update.message.reply_text(
        "👋 Привет, бро! Я бот для создания документов. 🚀\n"
        "Команды:\n"
        "/generate - Создать новый документ\n"
        "/list_saved - Показать сохранённые документы\n"
        "/cancel - Отменить текущую операцию\n"
        "Готов начать? Жми /generate! 😎"
    )

def start_generate(update: Update, context: CallbackContext) -> int:
    """Начало процесса генерации документа."""
    keyboard = [
        [InlineKeyboardButton("Imperative", callback_data='template_imperative')],
        [InlineKeyboardButton("UR Recruitment", callback_data='template_ur')],
        [InlineKeyboardButton("Small World", callback_data='template_small_world')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('📋 Выбери шаблон:', reply_markup=reply_markup)
    return SELECT_TEMPLATE

def template_selected(update: Update, context: CallbackContext) -> int:
    """Обработка выбора шаблона."""
    query = update.callback_query
    query.answer()
    template = query.data
    context.user_data['template'] = template
    query.edit_message_text(text=f"✅ Выбран шаблон: {template.replace('template_', '').title()}")
    query.message.reply_text("✍️ Введи имя клиента:")
    return INPUT_NAME

def name_input(update: Update, context: CallbackContext) -> int:
    """Обработка ввода имени клиента."""
    context.user_data['client_name'] = update.message.text.strip()
    keyboard = [
        [InlineKeyboardButton("Текущая дата и время", callback_data='current_date')],
        [InlineKeyboardButton("Ввести свою дату", callback_data='custom_date')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('📅 Хочешь текущую дату по Киеву или свою?', reply_markup=reply_markup)
    return CHOOSE_DATE

def date_chosen(update: Update, context: CallbackContext) -> int:
    """Обработка выбора даты."""
    query = update.callback_query
    query.answer()
    if query.data == 'current_date':
        now = datetime.now(pytz.utc).astimezone(kiev_tz)
        date_time = now.strftime("%d.%m.%Y %H:%M")
        context.user_data['date_time'] = date_time
        generate_and_send_document(update, context)
        ask_to_save(update, context)
        return ASK_SAVE
    else:
        query.message.reply_text("📅 Введи дату и время в формате ДД.ММ.ГГГГ ЧЧ:ММ:")
        return INPUT_CUSTOM_DATE

def input_custom_date(update: Update, context: CallbackContext) -> int:
    """Обработка пользовательской даты."""
    try:
        date_time = datetime.strptime(update.message.text.strip(), "%d.%m.%Y %H:%M")
        context.user_data['date_time'] = update.message.text.strip()
        generate_and_send_document(update, context)
        ask_to_save(update, context)
        return ASK_SAVE
    except ValueError:
        update.message.reply_text("❌ Неверный формат. Введи дату в формате ДД.ММ.ГГГГ ЧЧ:ММ:")
        return INPUT_CUSTOM_DATE

def ask_to_save(update: Update, context: CallbackContext):
    """Запрос на сохранение конфигурации."""
    keyboard = [
        [InlineKeyboardButton("💾 Сохранить", callback_data='save')],
        [InlineKeyboardButton("🚫 Не сохранять", callback_data='dont_save')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.effective_message.reply_text('Хочешь сохранить эту конфигурацию документа?', reply_markup=reply_markup)

def save_decision(update: Update, context: CallbackContext) -> int:
    """Обработка решения о сохранении."""
    query = update.callback_query
    query.answer()
    if query.data == 'save':
        user_id = query.from_user.id
        template = context.user_data['template']
        client_name = context.user_data['client_name']
        date_time = context.user_data['date_time']
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute("INSERT INTO saved_documents (user_id, template, client_name, date) VALUES (?, ?, ?, ?)",
                  (user_id, template, client_name, date_time))
        conn.commit()
        conn.close()
        query.edit_message_text(text="💾 Конфигурация сохранена! 🎉")
    else:
        query.edit_message_text(text="🚫 Конфигурация не сохранена.")
    query.message.reply_text("🔄 Хочешь создать ещё один документ? Жми /generate 😎")
    return ConversationHandler.END

def cancel(update: Update, context: CallbackContext) -> int:
    """Отмена диалога."""
    update.message.reply_text('❌ Операция отменена. Хочешь начать заново? Жми /generate 😎')
    return ConversationHandler.END

def list_saved(update: Update, context: CallbackContext):
    """Список сохранённых конфигураций."""
    user_id = update.effective_user.id
    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    c.execute("SELECT id, template, client_name, date FROM saved_documents WHERE user_id=?", (user_id,))
    rows = c.fetchall()
    conn.close()
    if not rows:
        update.message.reply_text("📭 У тебя нет сохранённых документов.")
    else:
        text = "📋 Твои сохранённые документы:\n"
        for row in rows:
            text += f"🆔 {row[0]} | Шаблон: {row[1].replace('template_', '').title()} | Клиент: {row[2]} | Дата: {row[3]}\n"
        update.message.reply_text(text)

@app.route('/webhook', methods=['POST'])
def webhook():
    """Обработка входящих обновлений Telegram."""
    try:
        update = Update.de_json(request.get_json(force=True), updater.bot)
        dispatcher.process_update(update)
    except Exception as e:
        print(f"Webhook error: {e}")
    return 'OK'

@app.route('/ping')
def ping():
    """Эндпоинт для Uptime Robot."""
    return 'OK'

if __name__ == '__main__':
    # Инициализация базы данных
    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS saved_documents
                 (id INTEGER PRIMARY KEY, user_id INTEGER, template TEXT, client_name TEXT, date TEXT)''')
    conn.commit()
    conn.close()

    # Создание шаблонов
    initialize_templates()

    # Настройка бота
    updater = Updater(token=TOKEN, use_context=True)
    dispatcher = updater.dispatcher

    # Обработчик диалога
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('generate', start_generate)],
        states={
            SELECT_TEMPLATE: [CallbackQueryHandler(template_selected)],
            INPUT_NAME: [MessageHandler(Filters.text & ~Filters.command, name_input)],
            CHOOSE_DATE: [CallbackQueryHandler(date_chosen)],
            INPUT_CUSTOM_DATE: [MessageHandler(Filters.text & ~Filters.command, input_custom_date)],
            ASK_SAVE: [CallbackQueryHandler(save_decision)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )

    # Добавление обработчиков
    dispatcher.add_handler(CommandHandler('start', start))
    dispatcher.add_handler(conv_handler)
    dispatcher.add_handler(CommandHandler('list_saved', list_saved))

    # Установка вебхука
    try:
        updater.bot.set_webhook(url=WEBHOOK_URL)
        print(f"Webhook установлен: {WEBHOOK_URL}")
    except Exception as e:
        print(f"Ошибка установки вебхука: {e}")

    # Запуск Flask
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
